#!/usr/bin/env python3
"""
Advanced Word Document Processor with Enhanced Find and Replace
Provides more sophisticated text processing and replacement capabilities
"""

import os
import re
import json
from typing import List, Dict, Tuple, Optional, Any
from pathlib import Path
import logging
from datetime import datetime
import subprocess
import tempfile
import shutil
import platform

try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("python-docx library not found. Install with: pip install python-docx")
    exit(1)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AdvancedWordProcessor:
    """Enhanced Word document processor with advanced find and replace capabilities"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
        self.backup_dir = Path("backups")
        self.backup_dir.mkdir(exist_ok=True)
    
    def _is_macos(self) -> bool:
        return platform.system() == 'Darwin'
    
    def _has_textutil(self) -> bool:
        try:
            subprocess.run(["which", "textutil"], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return True
        except Exception:
            return False
    
    def _convert_doc_to_docx(self, doc_path: str) -> Optional[str]:
        """Convert a .doc file to a temporary .docx using textutil on macOS."""
        if not self._is_macos() or not self._has_textutil():
            logger.warning(".doc conversion requires macOS textutil; skipping conversion.")
            return None
        try:
            temp_dir = tempfile.mkdtemp(prefix="doc_convert_")
            out_path = str(Path(temp_dir) / (Path(doc_path).stem + ".docx"))
            subprocess.run(["textutil", "-convert", "docx", doc_path, "-output", out_path], check=True)
            return out_path
        except Exception as e:
            logger.error(f"Failed to convert .doc to .docx for {doc_path}: {e}")
            return None
    
    def _convert_docx_to_doc(self, docx_path: str, dest_doc_path: str) -> bool:
        """Convert a .docx file back to .doc using textutil on macOS."""
        if not self._is_macos() or not self._has_textutil():
            logger.warning(".doc back-conversion requires macOS textutil; skipping conversion.")
            return False
        try:
            subprocess.run(["textutil", "-convert", "doc", docx_path, "-output", dest_doc_path], check=True)
            return True
        except Exception as e:
            logger.error(f"Failed to convert .docx back to .doc for {docx_path}: {e}")
            return False
    
    def is_word_file(self, file_path: str) -> bool:
        """Check if file is a supported Word document"""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def create_backup(self, file_path: str) -> str:
        """Create a backup of the original file"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_suffix = Path(file_path).suffix or ".docx"
        backup_name = f"{Path(file_path).stem}_{timestamp}{original_suffix}"
        backup_path = self.backup_dir / backup_name
        
        try:
            shutil.copy2(file_path, backup_path)
            logger.info(f"Created backup: {backup_path}")
            return str(backup_path)
        except Exception as e:
            logger.error(f"Failed to create backup: {e}")
            return ""
    
    def extract_text_with_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word document with structural information
        
        Returns:
            Dictionary containing:
            - paragraphs: List of paragraph texts with metadata
            - tables: List of table data
            - full_text: Complete document text
        """
        try:
            # If this is a .doc, convert to a temporary .docx for structured extraction
            working_path = file_path
            cleanup_paths: List[str] = []
            if Path(file_path).suffix.lower() == '.doc':
                converted = self._convert_doc_to_docx(file_path)
                if converted:
                    working_path = converted
                    cleanup_paths.append(converted)
                else:
                    # Cannot process .doc without conversion
                    return {'paragraphs': [], 'tables': [], 'full_text': '', 'file_path': file_path}
            
            doc = Document(working_path)
            result = {
                'paragraphs': [],
                'tables': [],
                'full_text': '',
                'file_path': file_path
            }
            
            # Extract paragraphs with metadata
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # Only include non-empty paragraphs
                    result['paragraphs'].append({
                        'index': i,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # Extract table data
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'index': table_idx,
                    'rows': []
                }
                for row_idx, row in enumerate(table.rows):
                    row_data = {
                        'index': row_idx,
                        'cells': []
                    }
                    for cell_idx, cell in enumerate(row.cells):
                        row_data['cells'].append({
                            'index': cell_idx,
                            'text': cell.text
                        })
                    table_data['rows'].append(row_data)
                result['tables'].append(table_data)
            
            # Create full text
            all_text_parts = [p['text'] for p in result['paragraphs']]
            for table in result['tables']:
                for row in table['rows']:
                    for cell in row['cells']:
                        all_text_parts.append(cell['text'])
            
            result['full_text'] = '\n'.join(all_text_parts)
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return {'paragraphs': [], 'tables': [], 'full_text': '', 'file_path': file_path}
        finally:
            # Clean up any temporary converted files
            try:
                if 'cleanup_paths' in locals():
                    for p in cleanup_paths:
                        # remove file and its temporary directory if empty
                        try:
                            Path(p).unlink(missing_ok=True)
                            temp_dir = Path(p).parent
                            # try removing directory if now empty
                            temp_dir.rmdir()
                        except Exception:
                            pass
            except Exception:
                pass
    
    def find_occurrences_with_context(
        self,
        file_path: str,
        search_term: str,
        context_chars: int = 150,
        case_sensitive: bool = False
    ) -> List[Dict]:
        """
        Find all occurrences with enhanced context and metadata
        
        Returns:
            List of occurrence dictionaries with detailed information
        """
        doc_structure = self.extract_text_with_structure(file_path)
        if not doc_structure['full_text']:
            return []
        
        occurrences = []
        flags = 0 if case_sensitive else re.IGNORECASE
        search_pattern = re.compile(re.escape(search_term), flags)
        
        for match in search_pattern.finditer(doc_structure['full_text']):
            start_pos = match.start()
            end_pos = match.end()
            
            # Get context around the match
            context_start = max(0, start_pos - context_chars)
            context_end = min(len(doc_structure['full_text']), end_pos + context_chars)
            
            context_before = doc_structure['full_text'][context_start:start_pos]
            context_after = doc_structure['full_text'][end_pos:context_end]
            full_context = doc_structure['full_text'][context_start:context_end]
            
            # Determine if match is in a table or paragraph
            location_type = "paragraph"
            location_index = 0
            
            # Check if match is in a table
            current_pos = 0
            for table_idx, table in enumerate(doc_structure['tables']):
                table_text = ""
                for row in table['rows']:
                    for cell in row['cells']:
                        cell_text = cell['text'] + "\n"
                        if current_pos <= start_pos < current_pos + len(cell_text):
                            location_type = "table"
                            location_index = table_idx
                            break
                        table_text += cell_text
                        current_pos += len(cell_text)
                    if location_type == "table":
                        break
                if location_type == "table":
                    break
            
            # If not in table, find paragraph
            if location_type == "paragraph":
                current_pos = 0
                for para_idx, paragraph in enumerate(doc_structure['paragraphs']):
                    para_text = paragraph['text'] + "\n"
                    if current_pos <= start_pos < current_pos + len(para_text):
                        location_index = para_idx
                        break
                    current_pos += len(para_text)
            
            occurrences.append({
                'file_path': file_path,
                'match_text': match.group(),
                'context_before': context_before,
                'context_after': context_after,
                'full_context': full_context,
                'context': full_context,
                'start_pos': start_pos,
                'end_pos': end_pos,
                'location_type': location_type,
                'location_index': location_index,
                'replacement_text': match.group().replace(search_term, search_term),  # Placeholder
                'unique_id': f"{Path(file_path).name}_{start_pos}_{end_pos}"
            })
        
        return occurrences
    
    def replace_text_advanced(self, file_path: str, old_text: str, new_text: str, 
                            occurrence_id: Optional[str] = None) -> Dict[str, Any]:
        """
        Advanced text replacement with specific occurrence targeting
        
        Args:
            file_path: Path to Word document
            old_text: Text to find
            new_text: Text to replace with
            occurrence_id: Specific occurrence ID to replace (if None, replaces all)
            
        Returns:
            Dictionary with replacement results
        """
        result = {
            'success': False,
            'replacements_made': 0,
            'backup_path': '',
            'error': None
        }
        
        try:
            # Create backup of the original file (preserve extension)
            backup_path = self.create_backup(file_path)
            result['backup_path'] = backup_path
            
            original_suffix = Path(file_path).suffix.lower()
            working_path = file_path
            temp_converted = None
            
            # If this is a .doc, convert to temp .docx for editing
            if original_suffix == '.doc':
                temp_converted = self._convert_doc_to_docx(file_path)
                if not temp_converted:
                    result['error'] = ".doc conversion unavailable on this system. Install textutil on macOS."
                    return result
                working_path = temp_converted
            
            doc = Document(working_path)
            replacements_made = 0
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    replacements_made += 1
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
                            replacements_made += 1
            
            if replacements_made > 0:
                # Save the edited document
                doc.save(working_path)
                
                if original_suffix == '.doc':
                    # Convert back to .doc, overwriting original
                    success = self._convert_docx_to_doc(working_path, file_path)
                    if not success:
                        result['error'] = "Failed to convert updated .docx back to .doc."
                        return result
                
                result['success'] = True
                result['replacements_made'] = replacements_made
                logger.info(f"Made {replacements_made} replacements in {file_path}")
            else:
                result['error'] = f"No occurrences of '{old_text}' found"
                logger.warning(result['error'])
                
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"Error replacing text in {file_path}: {e}")
        finally:
            # Cleanup temporary converted file
            try:
                if 'temp_converted' in locals() and temp_converted:
                    try:
                        Path(temp_converted).unlink(missing_ok=True)
                        Path(temp_converted).parent.rmdir()
                    except Exception:
                        pass
            except Exception:
                pass
        
        return result
    
    def scan_document_advanced(
        self,
        file_path: str,
        search_term: str,
        context_chars: int = 150,
        case_sensitive: bool = False
    ) -> Dict[str, Any]:
        """Scan a single document using the advanced processor."""
        result: Dict[str, Any] = {
            'success': False,
            'file_path': file_path,
            'occurrences': [],
            'error': None,
            'case_sensitive': case_sensitive
        }

        try:
            if not Path(file_path).exists():
                result['error'] = f"File {file_path} does not exist"
                return result

            if not self.is_word_file(file_path):
                result['error'] = "Unsupported file type"
                return result

            occurrences = self.find_occurrences_with_context(
                file_path,
                search_term,
                context_chars,
                case_sensitive=case_sensitive
            )
            result['occurrences'] = occurrences
            result['case_sensitive'] = case_sensitive
            result['success'] = True
            return result
        except Exception as exc:
            result['error'] = str(exc)
            logger.error(f"Error scanning document {file_path}: {exc}")
            result['case_sensitive'] = case_sensitive
            return result

    def scan_directory_advanced(
        self,
        directory_path: str,
        search_term: str,
        context_chars: int = 150,
        case_sensitive: bool = False
    ) -> Dict[str, Any]:
        """
        Advanced directory scanning with detailed results
        
        Returns:
            Dictionary containing scan results and statistics
        """
        directory = Path(directory_path)
        
        if not directory.exists():
            return {
                'success': False,
                'error': f"Directory {directory_path} does not exist",
                'files_scanned': 0,
                'files_processed': 0,
                'total_occurrences': 0,
                'occurrences': [],
                'case_sensitive': case_sensitive
            }
        
        # Find all Word files
        word_files = []
        for ext in self.supported_extensions:
            word_files.extend(directory.rglob(f"*{ext}"))
        
        all_occurrences = []
        files_with_matches = 0
        
        logger.info(f"Scanning {len(word_files)} Word files in {directory_path}")
        
        for file_path in word_files:
            if self.is_word_file(str(file_path)):
                occurrences = self.find_occurrences_with_context(
                    str(file_path),
                    search_term,
                    context_chars,
                    case_sensitive=case_sensitive
                )
                if occurrences:
                    files_with_matches += 1
                    all_occurrences.extend(occurrences)
                logger.info(f"Found {len(occurrences)} occurrences in {file_path}")
        
        return {
            'success': True,
            'files_scanned': len(word_files),
            'files_processed': len(word_files),
            'files_with_matches': files_with_matches,
            'total_occurrences': len(all_occurrences),
            'occurrences': all_occurrences,
            'search_term': search_term,
            'directory': str(directory),
            'case_sensitive': case_sensitive
        }
    
    def export_results(self, results: Dict[str, Any], output_file: str = "search_results.json"):
        """Export search results to JSON file"""
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            logger.info(f"Results exported to {output_file}")
        except Exception as e:
            logger.error(f"Failed to export results: {e}")

def main():
    """Test the AdvancedWordProcessor functionality"""
    processor = AdvancedWordProcessor()
    
    if len(os.sys.argv) < 3:
        print("Usage: python advanced_word_processor.py <directory> <search_term>")
        print("Example: python advanced_word_processor.py ./documents 'example text'")
        return
    
    directory = os.sys.argv[1]
    search_term = os.sys.argv[2]
    
    print(f"Advanced search for '{search_term}' in {directory}")
    results = processor.scan_directory_advanced(directory, search_term)
    
    if results['success']:
        print(f"\nScan Results:")
        print(f"Files scanned: {results['files_scanned']}")
        print(f"Files with matches: {results['files_with_matches']}")
        print(f"Total occurrences: {results['total_occurrences']}")
        
        # Export results
        processor.export_results(results)
        
        # Show first few occurrences
        for i, occurrence in enumerate(results['occurrences'][:5], 1):
            print(f"\n{i}. File: {Path(occurrence['file_path']).name}")
            print(f"   Context: ...{occurrence['context_before']}[{occurrence['match_text']}]{occurrence['context_after']}...")
    else:
        print(f"Error: {results['error']}")

if __name__ == "__main__":
    main()
