#!/usr/bin/env python3
"""
Word Document Processor for Global Find and Replace
Handles reading, searching, and modifying Microsoft Word documents
"""

import os
import re
from typing import List, Dict, Optional, Any
from pathlib import Path
import logging

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    print("python-docx library not found. Install with: pip install python-docx")
    exit(1)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WordProcessor:
    """Handles Word document operations for find and replace functionality"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
    
    def is_word_file(self, file_path: str) -> bool:
        """Check if file is a supported Word document"""
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract all text from a .docx file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def find_text_occurrences(self, file_path: str, search_term: str, context_chars: int = 100) -> List[Dict]:
        """
        Find all occurrences of search_term in the document with context
        
        Returns:
            List of dictionaries containing:
            - file_path: path to the file
            - paragraph_index: index of paragraph containing the match
            - match_text: the matched text
            - context_before: text before the match
            - context_after: text after the match
            - full_context: complete context around the match
        """
        text = self.extract_text_from_docx(file_path)
        if not text:
            return []
        
        occurrences = []
        search_pattern = re.compile(re.escape(search_term), re.IGNORECASE)
        
        for match in search_pattern.finditer(text):
            start_pos = match.start()
            end_pos = match.end()
            
            # Get context around the match
            context_start = max(0, start_pos - context_chars)
            context_end = min(len(text), end_pos + context_chars)
            
            context_before = text[context_start:start_pos]
            context_after = text[end_pos:context_end]
            # Find which paragraph this occurs in
            paragraph_index = text[:start_pos].count('\n')
            
            occurrences.append({
                'file_path': file_path,
                'paragraph_index': paragraph_index,
                'match_text': match.group(),
                'context_before': context_before,
                'context_after': context_after,
                'full_context': full_context,
                'context': full_context,
                'start_pos': start_pos,
                'end_pos': end_pos
            })
        
        return occurrences
    
    def replace_text_in_docx(self, file_path: str, old_text: str, new_text: str) -> bool:
        """Legacy convenience wrapper returning a boolean result."""
        return self.replace_text(file_path, old_text, new_text).get('success', False)
    
    def scan_document(self, file_path: str, search_term: str, context_chars: int = 100) -> Dict[str, Any]:
        """Scan a single document for occurrences of the search term."""
        result: Dict[str, Any] = {
            'success': False,
            'file_path': file_path,
            'occurrences': [],
            'error': None
        }

        try:
            if not Path(file_path).exists():
                result['error'] = f"File {file_path} does not exist"
                return result

            if not self.is_word_file(file_path):
                result['error'] = "Unsupported file type"
                return result

            occurrences = self.find_text_occurrences(file_path, search_term, context_chars)
            result['occurrences'] = occurrences
            result['success'] = True
            return result
        except Exception as exc:
            result['error'] = str(exc)
            logger.error(f"Error scanning document {file_path}: {exc}")
            return result

    def replace_text(self, file_path: str, old_text: str, new_text: str) -> Dict[str, Any]:
        """Replace text in a document and report the result."""
        result = {
            'success': False,
            'file_path': file_path,
            'replacements_made': 0,
            'error': None
        }

        try:
            if not Path(file_path).exists():
                result['error'] = f"File {file_path} does not exist"
                return result

            replacements_made = 0
            doc = Document(file_path)

            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    replacements_made += 1

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            cell.text = cell.text.replace(old_text, new_text)
                            replacements_made += 1

            if replacements_made > 0:
                doc.save(file_path)
                result['success'] = True
                result['replacements_made'] = replacements_made
            else:
                result['error'] = f"No occurrences of '{old_text}' found"
            return result
        except Exception as exc:
            logger.error(f"Error replacing text in {file_path}: {exc}")
            result['error'] = str(exc)
            return result

    def scan_directory(self, directory_path: str, search_term: str, context_chars: int = 100) -> Dict[str, Any]:
        """
        Scan a directory for Word files and find all occurrences of search_term
        
        Args:
            directory_path: Path to directory to scan
            search_term: Text to search for
            context_chars: Number of characters of context to include
            
        Returns:
            List of all occurrences found across all Word files
        """
        result: Dict[str, Any] = {
            'success': False,
            'directory': directory_path,
            'files_processed': 0,
            'total_occurrences': 0,
            'occurrences': [],
            'errors': []
        }

        all_occurrences: List[Dict[str, Any]] = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Directory {directory_path} does not exist")
            result['error'] = f"Directory {directory_path} does not exist"
            return result
        
        # Find all Word files in directory and subdirectories
        word_files = []
        for ext in self.supported_extensions:
            word_files.extend(directory.rglob(f"*{ext}"))
        
        logger.info(f"Found {len(word_files)} Word files in {directory_path}")
        
        for file_path in word_files:
            if self.is_word_file(str(file_path)):
                occurrences = self.find_text_occurrences(str(file_path), search_term, context_chars)
                all_occurrences.extend(occurrences)
                logger.info(f"Found {len(occurrences)} occurrences in {file_path}")
        
        result['success'] = True
        result['files_processed'] = len(word_files)
        result['total_occurrences'] = len(all_occurrences)
        result['occurrences'] = all_occurrences
        return result



def main():
    """Test the WordProcessor functionality"""
    processor = WordProcessor()
    
    # Example usage
    if len(os.sys.argv) < 3:
        print("Usage: python word_processor.py <directory> <search_term>")
        print("Example: python word_processor.py ./documents 'example text'")
        return
    
    directory = os.sys.argv[1]
    search_term = os.sys.argv[2]
    
    print(f"Searching for '{search_term}' in {directory}")
    occurrences = processor.scan_directory(directory, search_term)
    
    print(f"\nFound {len(occurrences)} occurrences:")
    for i, occurrence in enumerate(occurrences, 1):
        print(f"\n{i}. File: {occurrence['file_path']}")
        print(f"   Context: ...{occurrence['context_before']}[{occurrence['match_text']}]{occurrence['context_after']}...")

if __name__ == "__main__":
    main()
