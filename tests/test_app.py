#!/usr/bin/env python3
"""
Unit tests for Flask app.py
"""

import unittest
import tempfile
import os
import shutil
import json
from pathlib import Path
import sys

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app import app

class TestFlaskApp(unittest.TestCase):
    def setUp(self):
        """Set up test fixtures"""
        self.app = app.test_client()
        self.temp_dir = tempfile.mkdtemp()
        
        # Create test documents
        self.create_test_documents()
    
    def tearDown(self):
        """Clean up test fixtures"""
        shutil.rmtree(self.temp_dir)
    
    def create_test_documents(self):
        """Create test Word documents"""
        try:
            from docx import Document
            
            # Create a test document
            doc = Document()
            doc.add_heading('Test Document', 0)
            doc.add_paragraph('This is a test document with some sample text.')
            doc.add_paragraph('It contains the word "test" multiple times.')
            doc.add_paragraph('Another paragraph with test content.')
            
            test_doc_path = os.path.join(self.temp_dir, 'test_document.docx')
            doc.save(test_doc_path)
            
        except ImportError:
            # Skip tests if python-docx is not available
            self.skipTest("python-docx not available")
    
    def test_index_route(self):
        """Test the index route"""
        response = self.app.get('/')
        self.assertEqual(response.status_code, 200)
    
    def test_search_documents_success(self):
        """Test successful document search"""
        data = {
            'directory': self.temp_dir,
            'search_term': 'test',
            'context_chars': 50
        }
        
        response = self.app.post('/api/search', 
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 200)
        result = json.loads(response.data)
        self.assertIn('success', result)
    
    def test_search_documents_missing_directory(self):
        """Test search with missing directory"""
        data = {
            'search_term': 'test',
            'context_chars': 50
        }
        
        response = self.app.post('/api/search',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 400)
        result = json.loads(response.data)
        self.assertFalse(result['success'])
        self.assertIn('error', result)
    
    def test_search_documents_missing_search_term(self):
        """Test search with missing search term"""
        data = {
            'directory': self.temp_dir,
            'context_chars': 50
        }
        
        response = self.app.post('/api/search',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 400)
        result = json.loads(response.data)
        self.assertFalse(result['success'])
        self.assertIn('error', result)
    
    def test_search_documents_nonexistent_directory(self):
        """Test search with non-existent directory"""
        data = {
            'directory': '/nonexistent/directory',
            'search_term': 'test',
            'context_chars': 50
        }
        
        response = self.app.post('/api/search',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 400)
        result = json.loads(response.data)
        self.assertFalse(result['success'])
        self.assertIn('error', result)
    
    def test_replace_text_success(self):
        """Test successful text replacement"""
        data = {
            'file_path': os.path.join(self.temp_dir, 'test_document.docx'),
            'old_text': 'test',
            'new_text': 'example',
            'occurrence_id': 'test_1'
        }
        
        response = self.app.post('/api/replace',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 200)
        result = json.loads(response.data)
        self.assertIn('success', result)
    
    def test_replace_text_missing_parameters(self):
        """Test replace text with missing parameters"""
        data = {
            'file_path': os.path.join(self.temp_dir, 'test_document.docx'),
            'old_text': 'test'
            # Missing new_text
        }
        
        response = self.app.post('/api/replace',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 400)
        result = json.loads(response.data)
        self.assertFalse(result['success'])
        self.assertIn('error', result)
    
    def test_replace_all_success(self):
        """Test successful replace all"""
        data = {
            'occurrences': [
                {
                    'file_path': os.path.join(self.temp_dir, 'test_document.docx'),
                    'match_text': 'test',
                    'replacement_text': 'example',
                    'id': 'test_1'
                }
            ]
        }
        
        response = self.app.post('/api/replace_all',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 200)
        result = json.loads(response.data)
        self.assertTrue(result['success'])
        self.assertIn('total_processed', result)
        self.assertIn('successful_replacements', result)
    
    def test_replace_all_empty_occurrences(self):
        """Test replace all with empty occurrences"""
        data = {
            'occurrences': []
        }
        
        response = self.app.post('/api/replace_all',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 400)
        result = json.loads(response.data)
        self.assertFalse(result['success'])
        self.assertIn('error', result)
    
    def test_validate_directory_success(self):
        """Test successful directory validation"""
        data = {
            'directory': self.temp_dir
        }
        
        response = self.app.post('/api/validate_directory',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 200)
        result = json.loads(response.data)
        self.assertTrue(result['valid'])
        self.assertIn('word_files_count', result)
    
    def test_validate_directory_missing_path(self):
        """Test directory validation with missing path"""
        data = {}
        
        response = self.app.post('/api/validate_directory',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 200)
        result = json.loads(response.data)
        self.assertFalse(result['valid'])
        self.assertIn('error', result)
    
    def test_validate_directory_nonexistent(self):
        """Test directory validation with non-existent directory"""
        data = {
            'directory': '/nonexistent/directory'
        }
        
        response = self.app.post('/api/validate_directory',
                               data=json.dumps(data),
                               content_type='application/json')
        
        self.assertEqual(response.status_code, 200)
        result = json.loads(response.data)
        self.assertFalse(result['valid'])
        self.assertIn('error', result)
    
    def test_static_files_route(self):
        """Test static files route"""
        # This test assumes static files exist
        response = self.app.get('/static/style.css')
        # Should return 200 if file exists, 404 if not
        self.assertIn(response.status_code, [200, 404])

if __name__ == '__main__':
    unittest.main()
