#!/usr/bin/env python3
"""
Unit tests for advanced_word_processor.py
"""

import unittest
import tempfile
import os
import shutil
from pathlib import Path
import sys

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from advanced_word_processor import AdvancedWordProcessor

class TestAdvancedWordProcessor(unittest.TestCase):
    def setUp(self):
        """Set up test fixtures"""
        self.processor = AdvancedWordProcessor()
        self.temp_dir = tempfile.mkdtemp()
        
        # Create test documents
        self.create_test_documents()
    
    def tearDown(self):
        """Clean up test fixtures"""
        shutil.rmtree(self.temp_dir)
    
    def create_test_documents(self):
        """Create test Word documents"""
        try:
            from docx import Document
            
            # Create multiple test documents
            for i in range(3):
                doc = Document()
                doc.add_heading(f'Test Document {i+1}', 0)
                doc.add_paragraph(f'This is test document {i+1} with sample text.')
                doc.add_paragraph('It contains the word "test" multiple times.')
                doc.add_paragraph('Another paragraph with test content.')
                
                test_doc_path = os.path.join(self.temp_dir, f'test_document_{i+1}.docx')
                doc.save(test_doc_path)
            
            self.test_doc_path = os.path.join(self.temp_dir, 'test_document_1.docx')
            
        except ImportError:
            # Skip tests if python-docx is not available
            self.skipTest("python-docx not available")
    
    def test_supported_extensions(self):
        """Test that supported extensions are defined"""
        self.assertIsInstance(self.processor.supported_extensions, list)
        self.assertIn('.docx', self.processor.supported_extensions)
        self.assertIn('.doc', self.processor.supported_extensions)
    
    def test_scan_document_advanced(self):
        """Test advanced document scanning"""
        if not hasattr(self, 'test_doc_path'):
            self.skipTest("Test document not created")
        
        results = self.processor.scan_document_advanced(self.test_doc_path, 'test', 50)
        
        self.assertIsInstance(results, dict)
        self.assertIn('success', results)
        self.assertIn('occurrences', results)
        
        if results['success']:
            self.assertGreater(len(results['occurrences']), 0)
            # Check that each occurrence has required fields
            for occurrence in results['occurrences']:
                self.assertIn('file_path', occurrence)
                self.assertIn('match_text', occurrence)
                self.assertIn('context', occurrence)
                self.assertIn('start_pos', occurrence)
                self.assertIn('end_pos', occurrence)
    
    def test_scan_directory_advanced(self):
        """Test advanced directory scanning"""
        results = self.processor.scan_directory_advanced(self.temp_dir, 'test', 50)
        
        self.assertIsInstance(results, dict)
        self.assertIn('success', results)
        self.assertIn('files_processed', results)
        self.assertIn('total_occurrences', results)
        self.assertIn('occurrences', results)
        
        if results['success']:
            self.assertGreater(results['files_processed'], 0)
    
    def test_replace_text_advanced(self):
        """Test advanced text replacement"""
        if not hasattr(self, 'test_doc_path'):
            self.skipTest("Test document not created")
        
        # Create a copy for testing
        test_copy = os.path.join(self.temp_dir, 'test_copy.docx')
        shutil.copy2(self.test_doc_path, test_copy)
        
        # Replace 'test' with 'example'
        result = self.processor.replace_text_advanced(test_copy, 'test', 'example')
        
        self.assertIsInstance(result, dict)
        self.assertIn('success', result)
        self.assertIn('replacements_made', result)
    
    def test_replace_text_advanced_with_occurrence_id(self):
        """Test advanced text replacement with occurrence ID"""
        if not hasattr(self, 'test_doc_path'):
            self.skipTest("Test document not created")
        
        # Create a copy for testing
        test_copy = os.path.join(self.temp_dir, 'test_copy_with_id.docx')
        shutil.copy2(self.test_doc_path, test_copy)
        
        # Replace with occurrence ID
        result = self.processor.replace_text_advanced(
            test_copy, 'test', 'example', 'test_occurrence_1'
        )
        
        self.assertIsInstance(result, dict)
        self.assertIn('success', result)
    
    def test_scan_directory_advanced_nonexistent(self):
        """Test scanning non-existent directory"""
        results = self.processor.scan_directory_advanced('nonexistent_dir', 'test', 50)
        
        self.assertIsInstance(results, dict)
        self.assertFalse(results['success'])
    
    def test_context_extraction(self):
        """Test context extraction functionality"""
        if not hasattr(self, 'test_doc_path'):
            self.skipTest("Test document not created")
        
        results = self.processor.scan_document_advanced(self.test_doc_path, 'test', 20)
        
        if results['success'] and results['occurrences']:
            occurrence = results['occurrences'][0]
            self.assertIn('context', occurrence)
            self.assertIsInstance(occurrence['context'], str)
            self.assertGreater(len(occurrence['context']), 0)

    def test_case_sensitive_search(self):
        """Ensure case sensitivity flag alters search results"""
        if not hasattr(self, 'test_doc_path'):
            self.skipTest("Test document not created")

        insensitive = self.processor.scan_document_advanced(
            self.test_doc_path,
            'Test',
            50,
            case_sensitive=False
        )
        sensitive = self.processor.scan_document_advanced(
            self.test_doc_path,
            'Test',
            50,
            case_sensitive=True
        )

        if insensitive['success'] and sensitive['success']:
            self.assertGreater(len(insensitive['occurrences']), 0)
            self.assertGreater(len(insensitive['occurrences']), len(sensitive['occurrences']))
            self.assertTrue(
                any(match['match_text'].islower() for match in insensitive['occurrences']),
                "Insensitive search should capture lowercase variants"
            )

if __name__ == '__main__':
    unittest.main()
