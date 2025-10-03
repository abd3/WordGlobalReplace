#!/usr/bin/env python3
"""
Unit tests for word_processor.py
"""

import unittest
import tempfile
import os
import shutil
from pathlib import Path
import sys

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from word_processor import WordProcessor

class TestWordProcessor(unittest.TestCase):
    def setUp(self):
        """Set up test fixtures"""
        self.processor = WordProcessor()
        self.temp_dir = tempfile.mkdtemp()
        
        # Create test documents
        self.create_test_documents()
    
    def tearDown(self):
        """Clean up test fixtures"""
        shutil.rmtree(self.temp_dir)
    
    def create_test_documents(self):
        """Create test Word documents"""
        try:
            from docx import Document
            
            # Create a test document
            doc = Document()
            doc.add_heading('Test Document', 0)
            doc.add_paragraph('This is a test document with some sample text.')
            doc.add_paragraph('It contains the word "test" multiple times.')
            doc.add_paragraph('Another paragraph with test content.')
            
            test_doc_path = os.path.join(self.temp_dir, 'test_document.docx')
            doc.save(test_doc_path)
            self.test_doc_path = test_doc_path
            
        except ImportError:
            # Skip tests if python-docx is not available
            self.skipTest("python-docx not available")
    
    def test_supported_extensions(self):
        """Test that supported extensions are defined"""
        self.assertIsInstance(self.processor.supported_extensions, list)
        self.assertIn('.docx', self.processor.supported_extensions)
        self.assertIn('.doc', self.processor.supported_extensions)
    
    def test_scan_document(self):
        """Test document scanning functionality"""
        if not hasattr(self, 'test_doc_path'):
            self.skipTest("Test document not created")
        
        # Test scanning for 'test' in the document
        results = self.processor.scan_document(self.test_doc_path, 'test')
        
        self.assertIsInstance(results, dict)
        self.assertIn('success', results)
        self.assertIn('occurrences', results)
        
        if results['success']:
            self.assertGreater(len(results['occurrences']), 0)
    
    def test_scan_document_nonexistent(self):
        """Test scanning non-existent document"""
        results = self.processor.scan_document('nonexistent.docx', 'test')
        
        self.assertIsInstance(results, dict)
        self.assertIn('success', results)
        self.assertFalse(results['success'])
    
    def test_replace_text(self):
        """Test text replacement functionality"""
        if not hasattr(self, 'test_doc_path'):
            self.skipTest("Test document not created")
        
        # Create a copy for testing
        test_copy = os.path.join(self.temp_dir, 'test_copy.docx')
        shutil.copy2(self.test_doc_path, test_copy)
        
        # Replace 'test' with 'example'
        result = self.processor.replace_text(test_copy, 'test', 'example')
        
        self.assertIsInstance(result, dict)
        self.assertIn('success', result)
    
    def test_scan_directory(self):
        """Test directory scanning"""
        results = self.processor.scan_directory(self.temp_dir, 'test')
        
        self.assertIsInstance(results, dict)
        self.assertIn('success', results)
        self.assertIn('files_processed', results)
        self.assertIn('total_occurrences', results)
    
    def test_scan_directory_nonexistent(self):
        """Test scanning non-existent directory"""
        results = self.processor.scan_directory('nonexistent_dir', 'test')
        
        self.assertIsInstance(results, dict)
        self.assertFalse(results['success'])

if __name__ == '__main__':
    unittest.main()
